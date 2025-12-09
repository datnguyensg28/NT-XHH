from docx import Document
from docx.shared import Cm
import io

def build_docx(template_path, text_map, image_map):
    doc = Document(template_path)

    # ===== REPLACE TEXT =====
    for p in doc.paragraphs:
        for key, val in text_map.items():
            token = f"{{{{{key}}}}}"
            if token in p.text:
                p.clear()
                p.add_run(str(val))

    # ===== REPLACE IMAGE =====
    for p in doc.paragraphs:
        for key, img_bytes in image_map.items():
            token = f"{{{{{key}}}}}"
            if token in p.text:
                p.clear()
                run = p.add_run()
                run.add_picture(io.BytesIO(img_bytes), width=Cm(12))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
