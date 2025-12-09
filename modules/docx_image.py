# modules/docx_image.py
import io
import zipfile
import re
from docx import Document
from docx.shared import Cm

def _merge_xml(xml: str) -> str:
    """
    Merge các text node bị split trong DOCX:
    - ghép các </w:t> <w:t ...> liên tiếp
    - loại bỏ whitespace giữa nodes để placeholder không bị split
    """
    # cơ bản: xóa ranh giới giữa text nodes
    xml = re.sub(r"</w:t>\s*<w:t[^>]*>", "", xml)
    xml = re.sub(r"</w:t><w:t[^>]*>", "", xml)
    return xml

def replace_text_bytes(docx_bytes: bytes, placeholder: str, value: str) -> bytes:
    """
    Replace text trong DOCX file bytes.
    placeholder: exact string to replace (e.g. "$ngaybatdau" or "${ngaybatdau}")
    """
    bio = io.BytesIO(docx_bytes)

    with zipfile.ZipFile(bio, "r") as zin:
        files = {f.filename: zin.read(f.filename) for f in zin.infolist()}

    xml = files.get("word/document.xml").decode("utf-8")
    xml = _merge_xml(xml)

    # replace all occurrences (không phân biệt có dấu ngoặc hay không)
    xml = xml.replace(placeholder, value)

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w") as zout:
        for name, content in files.items():
            if name == "word/document.xml":
                zout.writestr(name, xml.encode("utf-8"))
            else:
                zout.writestr(name, content)

    return out.getvalue()

def insert_image_into_docx_bytes(docx_bytes: bytes, placeholder: str, img_bytes: bytes, width_cm: float = 10):
    """
    Chèn hình. Thực hiện 2 bước:
    1) Merge XML và try chèn bằng python-docx (tìm paragraph chứa placeholder bằng cách normalize).
    2) Nếu không tìm thấy, thực hiện replace trực tiếp trong document.xml bằng cách thay placeholder bằng
       một paragraph trống ghi rõ marker <!--IMG:{ph}--> để python-docx có thể đọc lại và chèn.
    """
    # --- 1) Merge xml và chuẩn bị files dict ---
    bio = io.BytesIO(docx_bytes)
    with zipfile.ZipFile(bio, "r") as zin:
        files = {f.filename: zin.read(f.filename) for f in zin.infolist()}

    xml = files["word/document.xml"].decode("utf-8")
    xml = _merge_xml(xml)
    files["word/document.xml"] = xml.encode("utf-8")

    # Ghi tạm docx merged
    merged = io.BytesIO()
    with zipfile.ZipFile(merged, "w") as zout:
        for name, content in files.items():
            zout.writestr(name, content)
    merged.seek(0)

    # --- 2) Try với python-docx (an toàn) ---
    doc = Document(merged)

    norm_ph = placeholder.replace(" ", "").replace("\n", "").replace("\t", "").lower()
    found = False

    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        norm_full = full.replace(" ", "").replace("\n", "").replace("\t", "").lower()

        if norm_ph in norm_full:
            # remove all runs
            for r in list(p.runs):
                try:
                    r._element.getparent().remove(r._element)
                except:
                    pass
            run = p.add_run()
            run.add_picture(io.BytesIO(img_bytes), width=Cm(width_cm))
            found = True

    if found:
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()

    # --- 3) fallback: sửa document.xml trực tiếp để đặt marker rồi chèn lại bằng python-docx ---
    # Tạo một paragraph xml nhỏ thay placeholder bằng <!--IMG:ph--> marker
    # Lưu ý: đây là một fallback đơn giản, không sinh drawing, nhưng python-docx sẽ thấy comment text và ta chèn sau.
    placeholder_patterns = [
        placeholder,
        placeholder.replace("${", "$").replace("}", ""),
        placeholder.replace("$", "${") + "}",
    ]

    xml_text = files["word/document.xml"].decode("utf-8")
    replaced = False
    for pat in set(placeholder_patterns):
        if pat in xml_text:
            # thay thành một đoạn rõ ràng (một paragraph chứa marker)
            xml_text = xml_text.replace(pat, f"<!--IMG:{placeholder}-->")
            replaced = True

    if replaced:
        files["word/document.xml"] = xml_text.encode("utf-8")
        temp = io.BytesIO()
        with zipfile.ZipFile(temp, "w") as zout:
            for name, content in files.items():
                zout.writestr(name, content)
        temp.seek(0)

        # read with python-docx and tìm marker
        doc2 = Document(temp)
        for p in doc2.paragraphs:
            if f"IMG:{placeholder}" in p.text:
                # remove runs then add image
                for r in list(p.runs):
                    try:
                        r._element.getparent().remove(r._element)
                    except:
                        pass
                run = p.add_run()
                run.add_picture(io.BytesIO(img_bytes), width=Cm(width_cm))
        out2 = io.BytesIO()
        doc2.save(out2)
        return out2.getvalue()

    # Nếu vẫn không tìm được - trả về nguyên bản (không thay)
    return docx_bytes
