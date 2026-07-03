import io
from docx import Document
from docx.shared import Cm

def load_docx_bytes(docx_bytes: bytes):
    return Document(io.BytesIO(docx_bytes))

def save_docx(doc: Document) -> bytes:
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _replace_in_paragraph(paragraph, placeholder, value):
    # Lấy toàn bộ text gốc
    runs = paragraph.runs
    if not runs:
        return False

    full_text = "".join(r.text for r in runs)
    if placeholder not in full_text:
        return False

    # tạo text mới sau khi replace
    new_text = full_text.replace(placeholder, value)

    # phân bổ text mới vào các run theo chiều dài run cũ
    original_lens = [len(r.text) for r in runs]
    pos = 0

    for i, r in enumerate(runs):
        take = original_lens[i]

        # run cuối lấy phần còn lại
        if i == len(runs) - 1:
            r.text = new_text[pos:]
            break

        r.text = new_text[pos:pos + take]
        pos += take

    return True


def replace_text(doc: Document, placeholder: str, value: str):
    replaced = False

    # paragraphs
    for p in doc.paragraphs:
        if _replace_in_paragraph(p, placeholder, value):
            replaced = True

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _replace_in_paragraph(p, placeholder, value):
                        replaced = True

    return replaced


def _insert_img_to_paragraph(paragraph, placeholder, img_bytes, width_cm):
    runs = paragraph.runs
    full_text = "".join(r.text for r in runs)

    if placeholder not in full_text:
        return False

    # clear all runs → chèn ảnh
    for r in runs:
        r.text = ""

    run = paragraph.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=Cm(width_cm))
    return True


def insert_image(doc: Document, placeholder: str, img_bytes: bytes, width_cm=12):
    inserted = False

    # paragraphs
    for p in doc.paragraphs:
        if _insert_img_to_paragraph(p, placeholder, img_bytes, width_cm):
            inserted = True

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _insert_img_to_paragraph(p, placeholder, img_bytes, width_cm):
                        inserted = True

    return inserted


def _clear_cell(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""


def _add_image_to_cell(cell, img_bytes: bytes, width_cm=12):
    _clear_cell(cell)
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=Cm(width_cm))


def insert_image_in_final_table(doc: Document, image_no: int, title: str, img_bytes: bytes, width_cm=12):
    """
    Chen anh vao bang hinh anh cuoi file. Neu template chua co dong anh tuong ung,
    tu dong them dong moi de tranh mat anh khi tao bien ban.
    """
    image_table = None
    for table in reversed(doc.tables):
        if not table.rows or len(table.columns) < 2:
            continue
        first_row = [cell.text.strip().lower() for cell in table.rows[0].cells[:2]]
        if "hình" in first_row[1] or "hinh" in first_row[1]:
            image_table = table
            break

    if image_table is None:
        image_table = doc.add_table(rows=1, cols=2)
        image_table.rows[0].cells[0].text = "TÊN HẠNG MỤC"
        image_table.rows[0].cells[1].text = "HÌNH ẢNH"

    needle_texts = {
        str(image_no),
        f"anh{image_no}",
        f"anh {image_no}",
        f"ảnh{image_no}",
        f"ảnh {image_no}",
        f"${{anh{image_no}}}",
        f"$anh{image_no}",
    }
    title_norm = title.strip().lower()

    target_row = None
    for row in image_table.rows[1:]:
        cells = row.cells
        row_title = cells[0].text.strip().lower()
        row_image = cells[1].text.strip().lower()
        if row_title == title_norm or any(needle in row_image for needle in needle_texts):
            target_row = row
            break

    if target_row is None:
        target_row = image_table.add_row()
        target_row.cells[0].text = title

    _add_image_to_cell(target_row.cells[1], img_bytes, width_cm)
    return True
